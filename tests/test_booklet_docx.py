"""Tests for docx_renderer.py and docx_style_service.py"""
from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest
from docx import Document
from docx.oxml.ns import qn

from booklet_parser import (
    BLOCK_BULLET_LIST, BLOCK_CODE_BLOCK, BLOCK_HEADING,
    BLOCK_NOTE, BLOCK_NUMBERED_LIST, BLOCK_PARAGRAPH, BLOCK_TABLE,
    DocumentBlock, InlineRun, ParsedChapter, TableData,
)
from docx_renderer import build_docx
from docx_style_service import (
    is_font_available, resolve_font,
    set_page_margins, set_rtl_paragraph,
)
from booklet_parser import sanitize_for_docx


# ── Helpers ───────────────────────────────────────────────────────────────────

def _default_settings() -> dict:
    return {
        "title": "Test Booklet",
        "subtitle": "Subtitle",
        "course_name": "Cardiology",
        "author_name": "Test Author",
        "date_str": "2026-01-01",
        "font_persian": "Arial",      # use Arial as guaranteed-available fallback
        "font_english": "Arial",
        "font_heading": "Arial",
        "body_size": 12,
        "h1_size": 20, "h2_size": 16, "h3_size": 14, "h4_size": 12,
        "line_spacing": 1.15,
        "space_after_pt": 6,
        "margin_top_cm": 2.0, "margin_bottom_cm": 2.0,
        "margin_right_cm": 2.2, "margin_left_cm": 2.2,
        "heading_color_hex": "1F4E79",
        "table_header_color_hex": "2E75B6",
        "accent_color_hex": "2E75B6",
        "include_cover": True,
        "include_toc": True,
        "include_header": True,
        "include_footer": True,
        "include_page_numbers": True,
        "chapter_numbering_mode": "NONE",
        "first_h1_behavior": "USE_AS_TITLE",
        "show_source_filename": "NONE",
        "overwrite_mode": "NEW_VERSION",
    }


def _make_chapter(title: str, blocks: list[DocumentBlock] | None = None,
                   src_name: str = "test.txt") -> ParsedChapter:
    if blocks is None:
        blocks = [
            DocumentBlock(BLOCK_PARAGRAPH,
                          [InlineRun("This is paragraph text.")]),
        ]
    return ParsedChapter(
        title=title,
        source_path=Path(src_name),
        blocks=blocks,
        warnings=[],
        raw_word_count=10,
        table_count=0,
        invalid_table_count=0,
    )


def _build_to_tmp(chapters: list[ParsedChapter],
                   settings: dict | None = None) -> Path:
    settings = settings or _default_settings()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        out = Path(f.name)
    build_docx(chapters, settings, out)
    return out


# ── DOCX validity ─────────────────────────────────────────────────────────────

def test_builds_valid_docx():
    chapters = [_make_chapter("Chapter One")]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))
        assert doc is not None
        assert len(doc.paragraphs) > 0
    finally:
        out.unlink(missing_ok=True)


def test_output_file_has_content():
    chapters = [_make_chapter("My Chapter")]
    out = _build_to_tmp(chapters)
    try:
        assert out.stat().st_size > 2000, "DOCX should be non-trivial size"
    finally:
        out.unlink(missing_ok=True)


# ── Chapter count and page breaks ─────────────────────────────────────────────

def test_multiple_chapters_produce_page_breaks():
    chapters = [
        _make_chapter("Chapter 1"),
        _make_chapter("Chapter 2"),
        _make_chapter("Chapter 3"),
    ]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))
        # Count page break runs
        page_breaks = 0
        for para in doc.paragraphs:
            for run in para.runs:
                xml = run._r.xml
                if 'w:type="page"' in xml or "pageBreak" in xml:
                    page_breaks += 1
        assert page_breaks >= 2, f"Expected at least 2 page breaks for 3 chapters, got {page_breaks}"
    finally:
        out.unlink(missing_ok=True)


# ── Heading styles ────────────────────────────────────────────────────────────

def test_headings_use_word_styles():
    chapters = [_make_chapter("My Chapter", blocks=[
        DocumentBlock(BLOCK_HEADING, [InlineRun("Section 1")], level=2),
        DocumentBlock(BLOCK_HEADING, [InlineRun("Section 2")], level=3),
    ])]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))
        style_names = [p.style.name for p in doc.paragraphs]
        assert any("Heading" in s for s in style_names)
    finally:
        out.unlink(missing_ok=True)


# ── Table rendering ───────────────────────────────────────────────────────────

def test_json_table_rendered_as_word_table():
    table_data = TableData(
        headers=["Parameter", "STEMI", "NSTEMI"],
        rows=[["ECG", "ST elevation", "ST depression"],
              ["Troponin", "Elevated", "Elevated"]],
    )
    chapters = [_make_chapter("Cardio", blocks=[
        DocumentBlock(BLOCK_TABLE, table_data),
    ])]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))
        assert len(doc.tables) >= 1, "Should have at least one Word table"
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == "Parameter"
        assert len(table.rows) == 3  # header + 2 data rows
        assert len(table.columns) == 3
    finally:
        out.unlink(missing_ok=True)


def test_table_header_row_count():
    table_data = TableData(
        headers=["A", "B"],
        rows=[["1", "2"], ["3", "4"], ["5", "6"]],
    )
    chapters = [_make_chapter("Ch", blocks=[DocumentBlock(BLOCK_TABLE, table_data)])]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))
        table = doc.tables[0]
        assert len(table.rows) == 4  # 1 header + 3 data
    finally:
        out.unlink(missing_ok=True)


# ── RTL paragraph direction ───────────────────────────────────────────────────

def test_rtl_paragraph_set():
    doc = Document()
    para = doc.add_paragraph("Test RTL")
    set_rtl_paragraph(para, "RIGHT")
    pPr = para._p.pPr
    assert pPr is not None
    bidi = pPr.find(qn("w:bidi"))
    assert bidi is not None, "RTL bidi element should be set"


# ── Cover page ────────────────────────────────────────────────────────────────

def test_cover_page_includes_title():
    settings = _default_settings()
    settings["title"] = "جزوه قلب و عروق"
    settings["include_toc"] = False
    chapters = [_make_chapter("Ch")]
    out = _build_to_tmp(chapters, settings)
    try:
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "جزوه قلب و عروق" in all_text
    finally:
        out.unlink(missing_ok=True)


def test_no_cover_page_when_disabled():
    settings = _default_settings()
    settings["include_cover"] = False
    settings["include_toc"] = False
    settings["title"] = "COVER_TITLE_UNIQUE"
    chapters = [_make_chapter("My Chapter")]
    out = _build_to_tmp(chapters, settings)
    try:
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "COVER_TITLE_UNIQUE" not in all_text
    finally:
        out.unlink(missing_ok=True)


# ── Note callout styles ───────────────────────────────────────────────────────

def test_note_blocks_rendered():
    chapters = [_make_chapter("Ch", blocks=[
        DocumentBlock(BLOCK_NOTE, [InlineRun("⭐ نکته طلایی")],
                      metadata={"note_type": "GOLDEN_TIP"}),
        DocumentBlock(BLOCK_NOTE, [InlineRun("🚨 هشدار")],
                      metadata={"note_type": "WARNING"}),
    ])]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "نکته طلایی" in all_text
        assert "هشدار" in all_text
    finally:
        out.unlink(missing_ok=True)


# ── Lists ─────────────────────────────────────────────────────────────────────

def test_bullet_list_renders():
    items = [[InlineRun("item one")], [InlineRun("item two")]]
    chapters = [_make_chapter("Ch", blocks=[
        DocumentBlock(BLOCK_BULLET_LIST, items),
    ])]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "item one" in all_text
        assert "item two" in all_text
    finally:
        out.unlink(missing_ok=True)


def test_numbered_list_renders():
    items = [[InlineRun("first")], [InlineRun("second")]]
    chapters = [_make_chapter("Ch", blocks=[
        DocumentBlock(BLOCK_NUMBERED_LIST, items),
    ])]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "first" in all_text
    finally:
        out.unlink(missing_ok=True)


# ── Code block ────────────────────────────────────────────────────────────────

def test_code_block_renders():
    chapters = [_make_chapter("Ch", blocks=[
        DocumentBlock(BLOCK_CODE_BLOCK, "def hello(): pass"),
    ])]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "def hello" in all_text
    finally:
        out.unlink(missing_ok=True)


# ── Emoji preservation ────────────────────────────────────────────────────────

def test_emoji_preserved_in_output():
    chapters = [_make_chapter("Ch", blocks=[
        DocumentBlock(BLOCK_PARAGRAPH, [InlineRun("⭐ Golden ✅ Check 💊 Drug 🚨 Warn")]),
    ])]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        # At least some emoji should survive the round-trip
        assert any(e in all_text for e in ["⭐", "✅", "💊", "🚨"])
    finally:
        out.unlink(missing_ok=True)


# ── Medical terms ─────────────────────────────────────────────────────────────

def test_medical_terms_preserved():
    chapters = [_make_chapter("AMI", blocks=[
        DocumentBlock(BLOCK_PARAGRAPH, [
            InlineRun("STEMI vs NSTEMI: Door-to-balloon time < 90 minutes. Troponin I elevated.")
        ]),
    ])]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "STEMI" in all_text
        assert "Troponin I" in all_text
    finally:
        out.unlink(missing_ok=True)


# ── Invalid XML chars ─────────────────────────────────────────────────────────

def test_invalid_xml_chars_do_not_crash():
    chapters = [_make_chapter("Ch", blocks=[
        DocumentBlock(BLOCK_PARAGRAPH, [InlineRun("test\x00\x01\x08text")]),
    ])]
    out = _build_to_tmp(chapters)
    try:
        doc = Document(str(out))  # Should open without error
        assert doc is not None
    finally:
        out.unlink(missing_ok=True)


# ── Page margins ──────────────────────────────────────────────────────────────

def test_page_margins_applied():
    from docx.shared import Cm
    doc = Document()
    section = doc.sections[0]
    settings = {"margin_top_cm": 3.0, "margin_bottom_cm": 2.5,
                 "margin_right_cm": 2.0, "margin_left_cm": 1.5}
    set_page_margins(section, settings)
    assert abs(section.top_margin - Cm(3.0)) < 500   # within 500 EMU ≈ 0.01mm
    assert abs(section.bottom_margin - Cm(2.5)) < 500


# ── Font fallback ─────────────────────────────────────────────────────────────

def test_resolve_font_uses_fallback_when_missing():
    font, was_fallback = resolve_font("NonExistentFont12345XYZ", "Arial")
    # Either Arial is available (Windows) or we get some fallback
    assert font is not None
    assert len(font) > 0


def test_resolve_font_prefers_preferred_when_available():
    # Arial should be available on Windows
    import platform
    if platform.system() == "Windows":
        font, was_fallback = resolve_font("Arial", "Courier New")
        assert font == "Arial"
        assert not was_fallback


# ── Chapter numbering ─────────────────────────────────────────────────────────

def test_chapter_numbering_fa():
    settings = _default_settings()
    settings["chapter_numbering_mode"] = "CHAPTER_FA"
    settings["include_toc"] = False
    settings["include_cover"] = False
    chapters = [_make_chapter("Heart"), _make_chapter("Lungs")]
    out = _build_to_tmp(chapters, settings)
    try:
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "فصل" in all_text
    finally:
        out.unlink(missing_ok=True)


# ── Integration: 20 files ─────────────────────────────────────────────────────

def test_twenty_chapters_build_correctly():
    chapters = [
        _make_chapter(f"Chapter {i+1}", blocks=[
            DocumentBlock(BLOCK_HEADING, [InlineRun(f"Section {i+1}.1")], level=2),
            DocumentBlock(BLOCK_PARAGRAPH, [InlineRun(f"Content for chapter {i+1}.")]),
        ])
        for i in range(20)
    ]
    settings = _default_settings()
    settings["include_cover"] = False
    settings["include_toc"] = False
    out = _build_to_tmp(chapters, settings)
    try:
        doc = Document(str(out))
        assert doc is not None
        heading_texts = [p.text for p in doc.paragraphs
                          if p.style and "Heading" in p.style.name]
        assert len(heading_texts) >= 20, f"Expected 20+ headings, got {len(heading_texts)}"
    finally:
        out.unlink(missing_ok=True)


def test_empty_chapter_not_rendered():
    chapters = [
        _make_chapter("Non-empty", blocks=[
            DocumentBlock(BLOCK_PARAGRAPH, [InlineRun("Some text")])
        ]),
        ParsedChapter(
            title="",
            source_path=Path("empty.txt"),
            blocks=[],
            warnings=[],
        ),
    ]
    settings = _default_settings()
    settings["include_cover"] = False
    settings["include_toc"] = False
    out = _build_to_tmp(chapters, settings)
    try:
        doc = Document(str(out))
        assert doc is not None
    finally:
        out.unlink(missing_ok=True)

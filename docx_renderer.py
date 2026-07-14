"""
docx_renderer.py – Build a professional DOCX booklet from ParsedChapter objects.

Responsibilities:
  - Cover page
  - Table of Contents (TOC) field
  - Chapter page-breaks and heading mapping
  - Paragraph / run rendering with RTL
  - Table rendering from TableData
  - Header, footer, and page number fields
  - Note callout styles
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Callable, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from booklet_parser import (
    BLOCK_BULLET_LIST, BLOCK_CODE_BLOCK, BLOCK_HEADING, BLOCK_HR,
    BLOCK_NOTE, BLOCK_NUMBERED_LIST, BLOCK_PARAGRAPH, BLOCK_QUOTE,
    BLOCK_TABLE,
    DocumentBlock, InlineRun, ParsedChapter, TableData,
    sanitize_for_docx,
)
from docx_style_service import (
    _get_or_add, _hex_to_rgb, _set_xml_bool,
    apply_font, get_note_style_name,
    set_page_margins, set_paragraph_spacing, set_rtl_paragraph, set_rtl_run,
    setup_document_styles, style_table,
)
from settings import (
    BOOKLET_ACCENT_COLOR_HEX, BOOKLET_DEFAULT_FONT_EMOJI,
    BOOKLET_DEFAULT_FONT_ENGLISH, BOOKLET_DEFAULT_FONT_PERSIAN,
    BOOKLET_HEADING_COLOR_HEX,
    BOOKLET_H1_FONT_SIZE_PT, BOOKLET_H2_FONT_SIZE_PT,
)


# ══════════════════════════════════════════════════════════════════════════════
# Word field helpers (PAGE, NUMPAGES, TOC)
# ══════════════════════════════════════════════════════════════════════════════

def _add_field(paragraph, field_code: str) -> None:
    """Insert a simple Word field (e.g. PAGE, NUMPAGES) into *paragraph*."""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = field_code

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)


def _add_page_number_field(paragraph, font_name: str, font_size: int,
                            label_before: str = "صفحه ", label_between: str = " از ") -> None:
    """Add 'صفحه X از Y' page number field."""
    if label_before:
        r = paragraph.add_run(label_before)
        r.font.name = font_name
        r.font.size = Pt(font_size)

    _add_field(paragraph, " PAGE ")

    if label_between:
        r2 = paragraph.add_run(label_between)
        r2.font.name = font_name
        r2.font.size = Pt(font_size)

    _add_field(paragraph, " NUMPAGES ")


def _insert_toc_field(doc: Document, font_persian: str) -> None:
    """Insert a TOC field paragraph (user must Update Field in Word)."""
    para = doc.add_paragraph()
    set_rtl_paragraph(para, "RIGHT")

    run = para.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    fldBegin.set(qn("w:dirty"), "true")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '

    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")

    run._r.append(fldBegin)
    run._r.append(instrText)
    run._r.append(fldEnd)


# ══════════════════════════════════════════════════════════════════════════════
# Cover page
# ══════════════════════════════════════════════════════════════════════════════

def _add_cover_page(doc: Document, settings: dict, font_persian: str, chapter_count: int) -> None:
    section = doc.sections[0]
    set_page_margins(section, settings)

    title_text = settings.get("title", "")
    subtitle_text = settings.get("subtitle", "")
    course_name = settings.get("course_name", "")
    university_name = settings.get("university_name", "")
    author_name = settings.get("author_name", "")
    ai_note = settings.get("ai_note", "")
    logo_path = settings.get("logo_path", "")
    date_str = settings.get("date_str", "")
    accent = settings.get("accent_color_hex", BOOKLET_ACCENT_COLOR_HEX)

    # Logo
    if logo_path:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(logo_path, width=Cm(5))
        except Exception:
            pass

    # Vertical space
    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Main title
    if title_text:
        para = doc.add_paragraph()
        set_rtl_paragraph(para, "CENTER")
        run = para.add_run(sanitize_for_docx(title_text))
        apply_font(run, font_persian, 28, bold=True, color_hex=accent)
        para.paragraph_format.space_after = Pt(8)

    # Subtitle
    if subtitle_text:
        para = doc.add_paragraph()
        set_rtl_paragraph(para, "CENTER")
        run = para.add_run(sanitize_for_docx(subtitle_text))
        apply_font(run, font_persian, 16, bold=False, color_hex="555555")
        para.paragraph_format.space_after = Pt(16)

    for _ in range(2):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Info lines
    info_lines = []
    if course_name:
        info_lines.append(f"درس: {course_name}")
    if university_name:
        info_lines.append(f"دانشگاه: {university_name}")
    if author_name:
        info_lines.append(f"تهیه‌کننده: {author_name}")
    if chapter_count:
        info_lines.append(f"تعداد فصل‌ها: {chapter_count}")
    if date_str:
        info_lines.append(f"تاریخ تولید: {date_str}")
    if ai_note:
        info_lines.append(ai_note)

    for line in info_lines:
        para = doc.add_paragraph()
        set_rtl_paragraph(para, "CENTER")
        run = para.add_run(sanitize_for_docx(line))
        apply_font(run, font_persian, 12)
        para.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# Header / Footer
# ══════════════════════════════════════════════════════════════════════════════

def _setup_header_footer(section, settings: dict, font_persian: str,
                          chapter_title: str = "") -> None:
    body_size = settings.get("body_size", 12)
    small_size = max(body_size - 2, 8)
    booklet_title = settings.get("title", "")
    accent = settings.get("accent_color_hex", BOOKLET_ACCENT_COLOR_HEX)

    if settings.get("include_header", True):
        header = section.header
        header.is_linked_to_previous = False
        # Clear existing paragraphs
        for p in header.paragraphs:
            p.clear()
        if header.paragraphs:
            hp = header.paragraphs[0]
        else:
            hp = header.add_paragraph()
        set_rtl_paragraph(hp, "RIGHT")
        header_text = booklet_title
        if chapter_title:
            header_text = f"{booklet_title} | {chapter_title}" if booklet_title else chapter_title
        run = hp.add_run(sanitize_for_docx(header_text))
        apply_font(run, font_persian, small_size, color_hex=accent)

    if settings.get("include_footer", True) and settings.get("include_page_numbers", True):
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()
        set_rtl_paragraph(fp, "CENTER")
        _add_page_number_field(fp, font_persian, small_size)


# ══════════════════════════════════════════════════════════════════════════════
# Inline runs → paragraph
# ══════════════════════════════════════════════════════════════════════════════

def _add_inline_runs(paragraph, runs: list, font_persian: str, font_english: str,
                     base_size: int, bold_override: bool = False) -> None:
    for run_obj in runs:
        if isinstance(run_obj, InlineRun):
            text = run_obj.text
            if not text:
                continue
            r = paragraph.add_run(sanitize_for_docx(text))
            is_bold = run_obj.bold or bold_override
            is_italic = run_obj.italic
            is_code = run_obj.code
            if is_code:
                apply_font(r, "Courier New", base_size - 1, bold=is_bold, italic=is_italic)
            else:
                apply_font(r, font_persian, base_size, bold=is_bold, italic=is_italic)
            set_rtl_run(r)
        elif isinstance(run_obj, str):
            r = paragraph.add_run(sanitize_for_docx(run_obj))
            apply_font(r, font_persian, base_size)
            set_rtl_run(r)


# ══════════════════════════════════════════════════════════════════════════════
# Block renderers
# ══════════════════════════════════════════════════════════════════════════════

def _render_heading(doc: Document, block: DocumentBlock, settings: dict,
                    font_heading: str, base_heading_level: int = 0) -> None:
    level = (block.level or 1) + base_heading_level
    level = min(max(level, 1), 4)
    style_name = f"Heading {level}"
    try:
        para = doc.add_paragraph(style=style_name)
    except Exception:
        para = doc.add_paragraph()

    set_rtl_paragraph(para, "RIGHT")
    heading_size = [0, BOOKLET_H1_FONT_SIZE_PT, BOOKLET_H2_FONT_SIZE_PT, 14, 12][level]
    content = block.content
    if isinstance(content, list):
        _add_inline_runs(para, content, font_heading, font_heading, heading_size, bold_override=True)
    else:
        r = para.add_run(sanitize_for_docx(str(content)))
        apply_font(r, font_heading, heading_size, bold=True,
                   color_hex=settings.get("heading_color_hex", BOOKLET_HEADING_COLOR_HEX))
        set_rtl_run(r)


def _render_paragraph(doc: Document, block: DocumentBlock, settings: dict,
                       font_persian: str, font_english: str) -> None:
    body_size = settings.get("body_size", 12)
    try:
        para = doc.add_paragraph(style="NLM Body Text")
    except Exception:
        para = doc.add_paragraph()
    set_rtl_paragraph(para, "JUSTIFY")
    set_paragraph_spacing(para)
    content = block.content
    if isinstance(content, list):
        _add_inline_runs(para, content, font_persian, font_english, body_size)
    else:
        r = para.add_run(sanitize_for_docx(str(content)))
        apply_font(r, font_persian, body_size)
        set_rtl_run(r)


def _render_note(doc: Document, block: DocumentBlock, settings: dict, font_persian: str) -> None:
    note_type = (block.metadata or {}).get("note_type", "IMPORTANT")
    style_name = get_note_style_name(note_type)
    body_size = settings.get("body_size", 12)
    try:
        para = doc.add_paragraph(style=style_name)
    except Exception:
        para = doc.add_paragraph()
    set_rtl_paragraph(para, "JUSTIFY")
    content = block.content
    if isinstance(content, list):
        _add_inline_runs(para, content, font_persian, font_persian, body_size)
    else:
        r = para.add_run(sanitize_for_docx(str(content)))
        apply_font(r, font_persian, body_size)
        set_rtl_run(r)


def _render_list(doc: Document, block: DocumentBlock, settings: dict,
                  font_persian: str, ordered: bool = False) -> None:
    body_size = settings.get("body_size", 12)
    list_style = "List Number" if ordered else "List Bullet"
    items = block.content if isinstance(block.content, list) else []
    for item in items:
        try:
            para = doc.add_paragraph(style=list_style)
        except Exception:
            para = doc.add_paragraph()
        set_rtl_paragraph(para, "RIGHT")
        if isinstance(item, list):
            _add_inline_runs(para, item, font_persian, font_persian, body_size)
        else:
            r = para.add_run(sanitize_for_docx(str(item)))
            apply_font(r, font_persian, body_size)
            set_rtl_run(r)


def _render_table(doc: Document, table_data: TableData, settings: dict,
                   font_persian: str) -> None:
    if not table_data.headers:
        return
    n_cols = len(table_data.headers)
    table = doc.add_table(rows=1 + len(table_data.rows), cols=n_cols)

    # Header row
    header_row = table.rows[0]
    for j, header_text in enumerate(table_data.headers):
        cell = header_row.cells[j]
        cell.text = ""
        para = cell.paragraphs[0]
        r = para.add_run(sanitize_for_docx(str(header_text)))
        apply_font(r, font_persian, settings.get("body_size", 12), bold=True)

    # Data rows
    for i, row_data in enumerate(table_data.rows):
        row = table.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = row.cells[j]
                cell.text = ""
                para = cell.paragraphs[0]
                r = para.add_run(sanitize_for_docx(str(cell_text)))
                apply_font(r, font_persian, settings.get("body_size", 12))

    style_table(table, settings, font_persian)

    # Spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _render_code_block(doc: Document, block: DocumentBlock, settings: dict) -> None:
    content = block.content if isinstance(block.content, str) else str(block.content)
    parse_error = (block.metadata or {}).get("parse_error", False)

    try:
        para = doc.add_paragraph(style="NLM Code")
    except Exception:
        para = doc.add_paragraph()

    para.paragraph_format.left_indent = Cm(0.5)
    if parse_error:
        r = para.add_run("[JSON Parse Error] ")
        r.font.color.rgb = RGBColor(0xCC, 0, 0)
        r.font.bold = True
        r.font.name = "Courier New"
    r2 = para.add_run(sanitize_for_docx(content))
    r2.font.name = "Courier New"
    r2.font.size = Pt(max(settings.get("body_size", 12) - 1, 9))


def _render_quote(doc: Document, block: DocumentBlock, settings: dict, font_persian: str) -> None:
    body_size = settings.get("body_size", 12)
    inner = block.content
    if isinstance(inner, list):
        for inner_block in inner:
            _render_block(doc, inner_block, settings, font_persian,
                          font_persian, font_persian, 0)
    else:
        para = doc.add_paragraph()
        set_rtl_paragraph(para, "JUSTIFY")
        para.paragraph_format.left_indent = Cm(1)
        r = para.add_run(sanitize_for_docx(str(inner)))
        apply_font(r, font_persian, body_size, italic=True)
        set_rtl_run(r)


def _render_hr(doc: Document) -> None:
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = _get_or_add(pPr, "w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    para.paragraph_format.space_after = Pt(4)


def _render_block(
    doc: Document, block: DocumentBlock, settings: dict,
    font_persian: str, font_english: str, font_heading: str,
    base_heading_level: int,
) -> None:
    t = block.block_type
    if t == BLOCK_HEADING:
        _render_heading(doc, block, settings, font_heading, base_heading_level)
    elif t == BLOCK_PARAGRAPH:
        _render_paragraph(doc, block, settings, font_persian, font_english)
    elif t == BLOCK_NOTE:
        _render_note(doc, block, settings, font_persian)
    elif t == BLOCK_BULLET_LIST:
        _render_list(doc, block, settings, font_persian, ordered=False)
    elif t == BLOCK_NUMBERED_LIST:
        _render_list(doc, block, settings, font_persian, ordered=True)
    elif t == BLOCK_TABLE:
        if isinstance(block.content, TableData):
            _render_table(doc, block.content, settings, font_persian)
    elif t == BLOCK_CODE_BLOCK:
        _render_code_block(doc, block, settings)
    elif t == BLOCK_QUOTE:
        _render_quote(doc, block, settings, font_persian)
    elif t == BLOCK_HR:
        _render_hr(doc)


# ══════════════════════════════════════════════════════════════════════════════
# Chapter numbering
# ══════════════════════════════════════════════════════════════════════════════

def _format_chapter_number(index: int, mode: str) -> str:
    """Return a chapter prefix string based on *mode*."""
    n = index + 1
    # Convert to Persian numerals
    fa_digits = "۰۱۲۳۴۵۶۷۸۹"
    fa_n = "".join(fa_digits[int(d)] for d in str(n))

    if mode == "NONE":
        return ""
    elif mode == "CHAPTER_FA":
        return f"فصل {fa_n}: "
    elif mode == "SECTION_FA":
        return f"بخش {fa_n}: "
    elif mode == "CHAPTER_EN":
        return f"Chapter {n}: "
    elif mode == "NUMBER_DOT":
        return f"{n}. "
    return ""


# ══════════════════════════════════════════════════════════════════════════════
# Source filename rendering
# ══════════════════════════════════════════════════════════════════════════════

def _add_source_caption(doc: Document, relative_path: str, font_persian: str,
                         body_size: int, location: str) -> None:
    """Add source filename caption (location = 'BELOW_TITLE' or 'END_OF_CHAPTER')."""
    caption_text = f"[منبع: {relative_path}]"
    para = doc.add_paragraph()
    set_rtl_paragraph(para, "RIGHT")
    r = para.add_run(sanitize_for_docx(caption_text))
    apply_font(r, font_persian, max(body_size - 2, 8), italic=True, color_hex="888888")
    set_rtl_run(r)


# ══════════════════════════════════════════════════════════════════════════════
# Chapter separator (page break)
# ══════════════════════════════════════════════════════════════════════════════

def _add_chapter_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)


# ══════════════════════════════════════════════════════════════════════════════
# Main render function
# ══════════════════════════════════════════════════════════════════════════════

def build_docx(
    chapters: list[ParsedChapter],
    settings: dict,
    output_path: Path,
    progress_callback: Optional[Callable[[int, int, str], None]] = None,
) -> dict:
    """Build a DOCX file from *chapters* and save to *output_path*.

    Returns a dict with build statistics.
    """
    doc = Document()

    # ── Setup page margins ────────────────────────────────────────────────────
    for section in doc.sections:
        set_page_margins(section, settings)

    # ── Setup styles ──────────────────────────────────────────────────────────
    font_info = setup_document_styles(doc, settings)
    font_persian = font_info["font_persian"]
    font_english = font_info["font_english"]
    font_heading = font_info["font_heading"]

    stats = {
        "chapter_count": 0,
        "word_count": 0,
        "table_count": 0,
        "invalid_table_count": 0,
        "font_warnings": [],
    }

    if font_info.get("font_persian_fallback"):
        stats["font_warnings"].append(
            f"Persian font '{settings.get('font_persian', '')}' is not installed. "
            f"Using '{font_persian}' instead."
        )
    if font_info.get("font_heading_fallback"):
        stats["font_warnings"].append(
            f"Heading font '{settings.get('font_heading', '')}' is not installed. "
            f"Using '{font_heading}' instead."
        )

    # ── Cover page ────────────────────────────────────────────────────────────
    if settings.get("include_cover", True):
        _add_cover_page(doc, settings, font_persian, len(chapters))

    # ── TOC ──────────────────────────────────────────────────────────────────
    if settings.get("include_toc", True):
        toc_title_para = doc.add_paragraph()
        set_rtl_paragraph(toc_title_para, "CENTER")
        toc_r = toc_title_para.add_run("فهرست مطالب")
        apply_font(toc_r, font_heading, BOOKLET_H1_FONT_SIZE_PT, bold=True,
                   color_hex=settings.get("heading_color_hex", BOOKLET_HEADING_COLOR_HEX))
        set_rtl_run(toc_r)
        toc_title_para.paragraph_format.space_after = Pt(8)

        _insert_toc_field(doc, font_persian)
        doc.add_page_break()

    # ── Setup header/footer on first section ──────────────────────────────────
    chapter_numbering_mode = settings.get("chapter_numbering_mode", "NONE")
    show_source_filename = settings.get("show_source_filename", "NONE")
    first_h1_behavior = settings.get("first_h1_behavior", "USE_AS_TITLE")
    body_size = settings.get("body_size", 12)

    # Header/footer on first content section
    if doc.sections:
        _setup_header_footer(doc.sections[0], settings, font_persian)

    # ── Chapters ──────────────────────────────────────────────────────────────
    active_chapters = [ch for ch in chapters if ch.blocks or ch.title]

    for ch_idx, chapter in enumerate(active_chapters):
        if progress_callback:
            progress_callback(ch_idx, len(active_chapters), chapter.title)

        # Page break between chapters (but not before the very first)
        if ch_idx > 0:
            _add_chapter_page_break(doc)

        # Chapter heading (Heading 1)
        prefix = _format_chapter_number(ch_idx, chapter_numbering_mode)
        chapter_title_text = prefix + chapter.title

        try:
            ch_para = doc.add_paragraph(style="Heading 1")
        except Exception:
            ch_para = doc.add_paragraph()
        set_rtl_paragraph(ch_para, "RIGHT")
        ch_r = ch_para.add_run(sanitize_for_docx(chapter_title_text))
        apply_font(ch_r, font_heading, BOOKLET_H1_FONT_SIZE_PT, bold=True,
                   color_hex=settings.get("heading_color_hex", BOOKLET_HEADING_COLOR_HEX))
        set_rtl_run(ch_r)
        ch_para.paragraph_format.keep_with_next = True

        # Source caption below title
        if show_source_filename == "BELOW_TITLE":
            _add_source_caption(doc, chapter.source_path.name, font_persian, body_size, "BELOW_TITLE")

        # Determine heading level offset (chapter is H1, so content starts at H2)
        # When first_h1_behavior == "USE_AS_TITLE" (default), we shift body headings by +1
        heading_offset = 1 if first_h1_behavior == "USE_AS_TITLE" else 0

        # Render blocks
        for block in chapter.blocks:
            _render_block(doc, block, settings, font_persian, font_english, font_heading, heading_offset)

        # Source caption at end
        if show_source_filename == "END_OF_CHAPTER":
            _add_source_caption(doc, chapter.source_path.name, font_persian, body_size, "END_OF_CHAPTER")

        stats["chapter_count"] += 1
        stats["word_count"] += chapter.raw_word_count
        stats["table_count"] += chapter.table_count
        stats["invalid_table_count"] += chapter.invalid_table_count

    # ── Save ─────────────────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return stats


def build_docx_bytes(
    chapters: list[ParsedChapter],
    settings: dict,
    progress_callback: Optional[Callable[[int, int, str], None]] = None,
) -> tuple[bytes, dict]:
    """Build DOCX in-memory and return (bytes, stats)."""
    import tempfile, os
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        tmp_path = Path(tf.name)
    try:
        stats = build_docx(chapters, settings, tmp_path, progress_callback)
        data = tmp_path.read_bytes()
    finally:
        tmp_path.unlink(missing_ok=True)
    return data, stats
